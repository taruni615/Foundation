#!/usr/bin/env python3
"""Render an exam-ready study-notes DOCX (and optionally PDF) from a notes JSON spec.

The layout reproduces the Foundation "Complete Exam-Ready Notes" booklets that live
in ``Generated_Notes/`` — coloured callout boxes, formula tables, derivation blocks,
periodic quick recaps and a chapter-end revision section.

Usage::

    python scripts/generate_notes_docx.py notes_spec/Class9_Physics.json
    python scripts/generate_notes_docx.py notes_spec/*.json --outdir Generated_Notes --pdf

The spec schema is documented in ``notes_spec/README.md``.
"""

from __future__ import annotations

import argparse
import glob
import json
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

TEMPLATE = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "notes_template.docx")

# ---------------------------------------------------------------- palette ----
INK = "1A1A1A"
NAVY = "1F3864"
BLUE = "2E5395"

# label colour, cell fill, border colour, border size (eighths of a point)
BOX = {
    "info": (NAVY, "EAF1FB", "6E8FC7", 8),
    "example": ("9C6318", "FDF3E7", "E8A33D", 8),
    "points": ("1E7B34", "EAF7EE", "6FBE8C", 8),
    "mistakes": ("D9534F", "FDEBEC", "D9534F", 8),
    "hook": ("8E44AD", "F3EAFB", "8E44AD", 8),
    "advanced": ("2C93A8", "E8F4F8", "2C93A8", 8),
    "result": (INK, "FFF9E6", "C9A227", 8),
    "recap": (INK, "F0F0F0", "999999", 8),
}

DIFFICULTY = {
    "easy": "\U0001f7e2 Easy   ★☆☆",
    "medium": "\U0001f7e1 Medium   ★★☆",
    "hard": "\U0001f534 Hard   ★★★",
}

RECAP_EVERY = 2  # a "Quick Recap" box is emitted after every N sub-topics


# ------------------------------------------------------------- primitives ----
def _rgb(hexcolor):
    return RGBColor.from_string(hexcolor)


def _set_style_sizes(doc):
    """Pin the heading sizes the booklet layout relies on."""
    for name, size in (("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 12)):
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = None


def para(doc, text="", *, style=None, bold=False, italic=False, size=None,
         color=None, align=None, space_before=None, space_after=None, indent=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold or None
        run.italic = italic or None
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = _rgb(color)
    pf = p.paragraph_format
    if align is not None:
        pf.alignment = align
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if indent is not None:
        pf.left_indent = Pt(indent)
    return p


def bullet(container, text, *, size=None):
    """A '●' bullet using the template's numbering definition (numId 1)."""
    p = container.add_paragraph(style="List Paragraph")
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), "1")
    numpr.append(ilvl)
    numpr.append(numid)
    p._p.get_or_add_pPr().append(numpr)
    run = p.add_run(text)
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p


def marked(container, marker, text, color, *, indent=11):
    """A '✓ '/'⚠ ' prefixed line inside a callout box."""
    p = container.add_paragraph()
    run = p.add_run(marker + "  ")
    run.bold = True
    run.font.color.rgb = _rgb(color)
    body = p.add_run(text)
    body.font.color.rgb = _rgb(INK)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(indent)
    return p


def _borders(table, color, size):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        inside = edge.startswith("inside")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4" if inside else str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color if not inside or size == 4 else "auto")
        borders.append(el)
    tbl_pr.append(borders)
    cell_mar = OxmlElement("w:tblCellMar")
    for edge, w in (("left", "10"), ("right", "10")):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:w"), w)
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tbl_pr.append(cell_mar)


def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _margins(cell, top, left, bottom, right):
    mar = OxmlElement("w:tcMar")
    for edge, w in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    cell._tc.get_or_add_tcPr().append(mar)


def _width_pct(cell, pct):
    tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        cell._tc.get_or_add_tcPr().append(tcw)
    tcw.set(qn("w:w"), str(pct))
    tcw.set(qn("w:type"), "pct")


def _full_width(table):
    tblw = table._tbl.tblPr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        table._tbl.tblPr.append(tblw)
    tblw.set(qn("w:w"), "5000")
    tblw.set(qn("w:type"), "pct")


def box(doc, kind, title=None, *, gap_after=7):
    """Create a single-cell coloured callout and return (cell, label_colour)."""
    label_color, fill, border, size = BOX[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _full_width(table)
    _borders(table, border, size)
    cell = table.rows[0].cells[0]
    _shade(cell, fill)
    _margins(cell, 140, 200, 140, 200)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    if title:
        p = cell.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = _rgb(label_color)
        p.paragraph_format.space_after = Pt(3)
    if gap_after:
        para(doc, space_after=gap_after)
    return cell, label_color


def body_text(cell, text, *, bold=False, space_after=None):
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold or None
    run.font.color.rgb = _rgb(INK)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


# ------------------------------------------------------------- components ----
def title_page(doc, spec):
    para(doc, spec["book_title"].upper(), bold=True, size=28, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=100)
    para(doc, spec.get("subtitle", "Complete Exam-Ready Notes"), italic=True, size=15,
         color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=15)
    para(doc, spec.get("tagline1", "Structured with Quick Summaries · Real-Life Examples · Formula Sheets"),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=30)
    para(doc, spec.get("tagline2", "Derivations · Memory Hooks · Common Mistakes · Chapter-End Revision"),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=5)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def meta_line(doc, topic):
    p = doc.add_paragraph()
    for label, value in (
        ("⏱ Estimated time: ", "%s    " % topic.get("time", "15 mins")),
        ("Difficulty: ", "%s    " % DIFFICULTY[topic.get("difficulty", "Medium").lower()]),
        ("Exam Importance: ", "⭐" * int(topic.get("importance", 4))),
    ):
        r = p.add_run(label)
        r.bold = True
        p.add_run(value)
    p.paragraph_format.space_after = Pt(7)


def detailed_explanation(doc, detail):
    para(doc, "\U0001f4d6 Detailed Explanation", style="Heading 3", color=NAVY, space_before=8)
    for key, label in (("definition", "Definition"), ("key_idea", "Key Idea"),
                       ("working_principle", "Working Principle")):
        if detail.get(key):
            para(doc, label, bold=True, italic=True, space_after=2)
            para(doc, detail[key], color=INK, space_after=6)
    if detail.get("applications"):
        para(doc, "Applications", bold=True, italic=True, space_after=2)
        for item in detail["applications"]:
            bullet(doc, item)
        para(doc, space_after=5)


def formula_table(doc, formula):
    para(doc, "\U0001f4d0 " + formula["name"], bold=True, color=BLUE,
         space_before=8, space_after=4)
    rows = [("Difficulty", DIFFICULTY[formula.get("difficulty", "Medium").lower()]),
            ("Formula", formula["formula"]),
            ("Meaning of Variables", formula.get("variables", "")),
            ("When to Use", formula.get("when_to_use", "")),
            ("Common Mistakes", formula.get("common_mistakes", "")),
            ("Shortcut / Memory Trick", formula.get("shortcut", ""))]
    rows = [r for r in rows if r[1]]
    table = doc.add_table(rows=len(rows), cols=2)
    _full_width(table)
    _borders(table, "C9A227", 4)
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        _shade(left, "F2E9C9")
        _shade(right, "FFF9E6")
        _width_pct(left, 1400)
        _width_pct(right, 3600)
        for c in (left, right):
            _margins(c, 90, 140, 90, 140)
        left.paragraphs[0].add_run(label).bold = True
        right.paragraphs[0].add_run(value)
    para(doc, space_after=8)


def derivation_block(doc, deriv):
    para(doc, "\U0001f9ee Derivation: " + deriv["title"], bold=True, color=BLUE,
         space_before=8, space_after=4)
    if deriv.get("why"):
        para(doc, "Why This Derivation Matters", bold=True, space_after=3)
        para(doc, deriv["why"], color=INK, space_after=6)
    if deriv.get("assumptions"):
        para(doc, "Assumptions Used", bold=True, space_after=3)
        for item in deriv["assumptions"]:
            bullet(doc, item)
    para(doc, "Step-by-Step Derivation", bold=True, space_after=3)
    for i, step in enumerate(deriv["steps"], 1):
        p = doc.add_paragraph()
        p.add_run("Step %d: " % i).bold = True
        p.add_run(step)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Pt(10)
    para(doc, "Final Result", bold=True, space_after=3)
    cell, _ = box(doc, "result", gap_after=4)
    body_text(cell, deriv["result"], bold=True)
    if deriv.get("exam_perspective"):
        para(doc, "Exam Perspective", bold=True, space_before=5, space_after=3)
        para(doc, deriv["exam_perspective"], color=INK, space_after=6)
    para(doc, space_after=8)


def quick_recap(doc, topics):
    para(doc, "⏸ Quick Recap — Topics Covered So Far", bold=True, size=10,
         color=INK, space_before=5, space_after=3)
    cell, _ = box(doc, "recap", gap_after=10)
    for t in topics:
        bullet(cell, "%s — %s" % (t["title"], t["quick_summary"]))


def sub_topic(doc, chapter_no, index, topic):
    para(doc, "%d.%d  %s" % (chapter_no, index, topic["title"]),
         style="Heading 2", color=BLUE, space_before=12)
    meta_line(doc, topic)

    cell, _ = box(doc, "info", "\U0001f50e Quick Summary")
    body_text(cell, topic["quick_summary"])

    if topic.get("real_life"):
        cell, _ = box(doc, "example", "\U0001f30d Real-Life Example")
        body_text(cell, topic["real_life"])

    if topic.get("detailed"):
        detailed_explanation(doc, topic["detailed"])

    if topic.get("points"):
        cell, color = box(doc, "points", "✓ Points to Remember (Exam Focus)")
        for item in topic["points"]:
            marked(cell, "✓", item, color)

    if topic.get("mistakes"):
        cell, color = box(doc, "mistakes", "⚠ Common Student Mistakes")
        for item in topic["mistakes"]:
            marked(cell, "⚠", item, color)

    if topic.get("memory_hook"):
        cell, _ = box(doc, "hook", "\U0001f9e0 Memory Hook")
        body_text(cell, topic["memory_hook"], bold=True)

    if topic.get("advanced"):
        cell, _ = box(doc, "advanced", "\U0001f680 Advanced Insights")
        body_text(cell, topic["advanced"])

    if topic.get("formulas"):
        para(doc, "Formula Section", style="Heading 3", color=NAVY, space_before=6)
        for formula in topic["formulas"]:
            formula_table(doc, formula)

    if topic.get("derivations"):
        para(doc, "Derivation Section", style="Heading 3", color=NAVY, space_before=6)
        for deriv in topic["derivations"]:
            derivation_block(doc, deriv)


def end_of_chapter(doc, chapter):
    topics = chapter["topics"]
    para(doc, "Chapter %d — End of Chapter Revision" % chapter["number"],
         style="Heading 1", color=NAVY, space_after=6)

    hooks = [t for t in topics if t.get("memory_hook")]
    if hooks:
        para(doc, "\U0001f9e0 All Memory Hooks", style="Heading 2", color="8E44AD", space_before=12)
        cell, _ = box(doc, "hook", gap_after=10)
        for t in hooks:
            bullet(cell, "%s: %s" % (t["title"], t["memory_hook"]))

    mistakes = [m for t in topics for m in t.get("mistakes", [])]
    if mistakes:
        para(doc, "⚠ All Common Mistakes", style="Heading 2", color="D9534F", space_before=12)
        cell, color = box(doc, "mistakes", gap_after=10)
        for m in mistakes:
            marked(cell, "⚠", m, color)

    formulas = [f for t in topics for f in t.get("formulas", [])]
    if formulas:
        para(doc, "\U0001f4d0 Formula Sheet", style="Heading 2", color=BLUE, space_before=12)
        table = doc.add_table(rows=len(formulas) + 1, cols=2)
        _full_width(table)
        _borders(table, "999999", 4)
        head = table.rows[0].cells
        for c, label in zip(head, ("Formula", "Meaning / Use")):
            _shade(c, BLUE)
            _margins(c, 90, 140, 90, 140)
            run = c.paragraphs[0].add_run(label)
            run.bold = True
            run.font.color.rgb = _rgb("FFFFFF")
        for i, f in enumerate(formulas, 1):
            left, right = table.rows[i].cells
            for c in (left, right):
                _margins(c, 90, 140, 90, 140)
            left.paragraphs[0].add_run("%s: " % f["name"]).bold = True
            left.paragraphs[0].add_run(f["formula"])
            right.paragraphs[0].add_run(f.get("when_to_use", ""))
        para(doc, space_after=10)

    para(doc, "⚡ Quick Revision Section", style="Heading 2", color=NAVY, space_before=12)
    cell, _ = box(doc, "info", gap_after=10)
    for t in topics:
        bullet(cell, "%s: %s" % (t["title"], t["quick_summary"]))


def chapter_block(doc, chapter):
    para(doc, "Chapter %d: %s" % (chapter["number"], chapter["title"]),
         style="Heading 1", color=NAVY, space_after=6)

    if chapter.get("prerequisites"):
        para(doc, "Prerequisites", style="Heading 2", color=NAVY, space_before=12)
        para(doc, "Before starting this chapter, make sure you are comfortable with the following ideas:",
             italic=True, color=INK, space_after=6)
        cell, _ = box(doc, "info", gap_after=12)
        for item in chapter["prerequisites"]:
            bullet(cell, item)

    topics = chapter["topics"]
    for i, topic in enumerate(topics, 1):
        sub_topic(doc, chapter["number"], i, topic)
        if i % RECAP_EVERY == 0 or i == len(topics):
            start = ((i - 1) // RECAP_EVERY) * RECAP_EVERY
            quick_recap(doc, topics[start:i])

    page_break(doc)
    end_of_chapter(doc, chapter)


# ------------------------------------------------------------------ build ----
def build(spec, out_path):
    doc = Document(TEMPLATE)
    _set_style_sizes(doc)

    header_text = spec.get("header", spec["book_title"] + " — Notes")
    for section in doc.sections:
        for p in section.header.paragraphs:
            for r in p.runs:
                r.text = header_text if "{{HEADER}}" in r.text else ""

    title_page(doc, spec)
    for chapter in spec["chapters"]:
        page_break(doc)
        chapter_block(doc, chapter)

    doc.core_properties.title = header_text
    doc.save(out_path)
    return out_path


def to_pdf(docx_path):
    """Export via Microsoft Word (the tool the reference booklets were produced with)."""
    import win32com.client

    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.Fields.Update()  # refresh the "Page X of Y" footer field
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)
    finally:
        word.Quit()
    return pdf_path


def load_spec(path):
    """Load a book spec, pulling in per-chapter files when ``chapters_dir`` is set.

    Splitting a book across ``<dir>/ch01.json``, ``ch02.json``, … keeps each chapter
    editable on its own; chapters are ordered by their ``number`` field.
    """
    with open(path, encoding="utf-8") as fh:
        spec = json.load(fh)

    if spec.get("chapters_dir"):
        base = os.path.join(os.path.dirname(os.path.abspath(path)), spec["chapters_dir"])
        chapters = []
        for chapter_path in sorted(glob.glob(os.path.join(base, "*.json"))):
            with open(chapter_path, encoding="utf-8") as fh:
                chapters.append(json.load(fh))
        spec["chapters"] = sorted(chapters, key=lambda c: c["number"])

    if not spec.get("chapters"):
        raise SystemExit("%s: no chapters found" % path)
    return spec


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("specs", nargs="+", help="notes spec JSON file(s); globs allowed")
    ap.add_argument("--outdir", default="Generated_Notes")
    ap.add_argument("--pdf", action="store_true", help="also export a PDF via Microsoft Word")
    args = ap.parse_args(argv)

    paths = [p for pattern in args.specs for p in (glob.glob(pattern) or [pattern])]
    os.makedirs(args.outdir, exist_ok=True)

    for path in paths:
        spec = load_spec(path)
        name = spec.get("output_name") or os.path.splitext(os.path.basename(path))[0]
        out = build(spec, os.path.join(args.outdir, name + ".docx"))
        print("wrote", out)
        if args.pdf:
            print("wrote", to_pdf(out))
    return 0


if __name__ == "__main__":
    sys.exit(main())
